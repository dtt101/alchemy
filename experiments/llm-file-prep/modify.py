
#!/usr/bin/env python3
# Reference implementation: extract segments from DOCX with stable IDs,
# validate an LLM-style patch JSON, and apply safe text edits back to the source
# while keeping paragraph structure (and most styling) intact.
#
# This script:
# 1) Loads the user's /mnt/data/test.docx
# 2) Extracts addressable paragraph segments with stable IDs (p-0001, ...)
# 3) Writes an ID map JSON for round-tripping
# 4) Prints a summary of the extracted segments
# 5) Applies a sample patch (replace_text, insert_after) using guardrails
# 6) Saves /mnt/data/test_patched.docx
#
# Notes:
# - For production, you’d anchor to bookmarks/content controls for true stability.
# - The patcher below preserves unaffected runs; for replaced spans, it rebuilds
#   the affected paragraph text and reuses the first run’s style for the new text.
# - This is a compact, readable baseline you can extend (e.g., to handle tables, shapes).

import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.oxml import OxmlElement

# ---------- Paths ----------
HERE = Path(__file__).resolve().parent
DOC_PATH = HERE / "test.docx"
OUT_MAP_JSON = HERE / "test.docx.map.json"
OUT_PATCHED = HERE / "test_patched.docx"

assert DOC_PATH.exists(), f"Missing {DOC_PATH}"

# ---------- Example patch (as if from an LLM) ----------

PATCH: Dict[str, Any] = {
    "doc_id": DOC_PATH.name,
    "edits": [
        {
            "id": "p-0003",
            "op": "replace_text",
            "old_excerpt": "This is an example document which I made to see how good mammoth is",
            "new_text": "This is an example document created to assess how well our extraction works."
        },
        {
            "id": "p-0005",
            "op": "insert_after",
            "new_text": "There should be a table immediately after this paragraph (note added by patcher)."
        }
    ]
}

# ---------- Data model ----------

@dataclass
class Segment:
    id: str
    kind: str  # "paragraph"
    para_idx: int
    text: str
    context_left: str
    context_right: str

# ---------- Extraction ----------

def enumerate_paragraphs(doc: Document) -> List[Paragraph]:
    # Document.paragraphs already gives block-level paragraphs (tables have their own cell paragraphs)
    return list(doc.paragraphs)


def make_segments(doc: Document, window: int = 80) -> List[Segment]:
    paras = enumerate_paragraphs(doc)
    segments: List[Segment] = []

    for i, p in enumerate(paras):
        full_text = p.text or ""
        left = full_text[:window]
        right = full_text[-window:]
        seg = Segment(
            id=f"p-{i+1:04d}",
            kind="paragraph",
            para_idx=i,
            text=full_text,
            context_left=left,
            context_right=right,
        )
        segments.append(seg)
    return segments

# ---------- Editing primitives ----------

def _remove_all_runs(paragraph: Paragraph) -> None:
    """Remove all run elements from a paragraph safely."""
    runs = list(paragraph.runs)
    for r in runs:
        r_el = r._element
        r_el.getparent().remove(r_el)


def replace_first_occurrence_in_runs(
    p: Paragraph, old_excerpt: str, new_text: str
) -> bool:
    """
    Replace the first occurrence of old_excerpt across a paragraph's runs with new_text.
    Rebuilds the paragraph *only for the affected span*, preserving unaffected runs,
    and copies the style of the first affected run.
    """
    if not old_excerpt:
        print("old_excerpt_not_provided")
        return False

    runs = list(p.runs)
    concatenated = "".join(r.text for r in runs)
    idx = concatenated.find(old_excerpt)
    if idx < 0:
        print("old_excerpt_not_found")
        return False

    start, end = idx, idx + len(old_excerpt)

    # Find first affected run for style reuse
    first_run_idx: Optional[int] = None
    pos = 0
    for ri, r in enumerate(runs):
        nxt = pos + len(r.text)
        if nxt > start and pos < end and first_run_idx is None:
            first_run_idx = ri
        pos = nxt
    if first_run_idx is None:
        first_run_idx = 0

    new_para_text = concatenated[:start] + new_text + concatenated[end:]

    # Clear and rebuild
    _remove_all_runs(p)
    new_run: Run = p.add_run(new_para_text)

    # Copy style (best-effort)
    src = runs[first_run_idx] if runs else None
    if src is not None:
        try:
            if src.style:
                new_run.style = src.style
            new_run.bold = src.bold
            new_run.italic = src.italic
            new_run.underline = src.underline
        except Exception:
            pass

    return True


def insert_paragraph_after(paragraph: Paragraph, text: str, copy_style: bool = True) -> Paragraph:
    """
    Insert a new paragraph *after* the given one, with optional style copy
    from the given paragraph's first run.
    """
    p_el = paragraph._element  # CT_P
    new_p_el = OxmlElement("w:p")
    p_el.addnext(new_p_el)
    new_para = Paragraph(new_p_el, paragraph._parent)

    if copy_style and paragraph.runs:
        src = paragraph.runs[0]
        run = new_para.add_run(text)
        try:
            if src.style:
                run.style = src.style
            run.bold = src.bold
            run.italic = src.italic
            run.underline = src.underline
        except Exception:
            pass
    else:
        new_para.add_run(text)

    return new_para

# ---------- Patch application ----------

def apply_patch_to_doc(doc: Document, segments: List[Segment], patch: Dict[str, Any]) -> Dict[str, Any]:
    """
    Apply a simple patch to the document using previously-extracted segments.
    Supported ops:
      - replace_text: requires {id, old_excerpt, new_text}
      - insert_after: requires {id, new_text}
    """
    report = {"applied": [], "skipped": []}
    seg_by_id = {s.id: s for s in segments}

    for edit in patch.get("edits", []):
        seg_id = edit.get("id")
        op = edit.get("op")

        if not seg_id or seg_id not in seg_by_id:
            report["skipped"].append({"id": seg_id, "reason": "unknown_id"})
            continue

        seg = seg_by_id[seg_id]
        try:
            p = doc.paragraphs[seg.para_idx]
        except IndexError:
            report["skipped"].append({"id": seg_id, "reason": "para_idx_out_of_range"})
            continue

        if op == "replace_text":
            old_excerpt = edit.get("old_excerpt", "")
            new_text = edit.get("new_text", "")
            if not new_text:
                report["skipped"].append({"id": seg_id, "reason": "empty_new_text"})
                continue
            ok = replace_first_occurrence_in_runs(p, old_excerpt, new_text)
            if ok:
                report["applied"].append(edit)
            else:
                report["skipped"].append({"id": seg_id, "reason": "old_excerpt_not_found"})

        elif op == "insert_after":
            new_text = edit.get("new_text", "")
            if not new_text:
                report["skipped"].append({"id": seg_id, "reason": "empty_new_text"})
                continue
            insert_paragraph_after(p, new_text, copy_style=True)
            report["applied"].append(edit)

        else:
            report["skipped"].append({"id": seg_id, "reason": f"unsupported_op:{op}"})

    return report

# ---------- Main flow ----------

def main() -> None:
    doc = Document(str(DOC_PATH))
    segments = make_segments(doc, window=80)

    # Save ID map
    id_map = [
        {
            "id": s.id,
            "kind": s.kind,
            "para_idx": s.para_idx,
            "text": s.text,
            "context_left": s.context_left,
            "context_right": s.context_right,
        }
        for s in segments
    ]
    with OUT_MAP_JSON.open("w", encoding="utf-8") as f:
        json.dump({"source": DOC_PATH.name, "segments": id_map}, f, ensure_ascii=False, indent=2)

    # Apply patch
    report = apply_patch_to_doc(doc, segments, PATCH)
    doc.save(str(OUT_PATCHED))

    print("== Extraction ==")
    print(f" Source: {DOC_PATH}")
    print(f" Segments: {len(segments)}")
    print(f" Map: {OUT_MAP_JSON}")
    print("\n== Patch Report ==")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    print(f"\n Output: {OUT_PATCHED}")


if __name__ == "__main__":
    main()
